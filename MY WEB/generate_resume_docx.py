from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_spacing(paragraph, space_before=0, space_after=0):
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)

def add_line_break(doc):
    p = doc.add_paragraph()
    set_spacing(p, 0, 0)

def add_section_header(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Calibri'
    set_spacing(p, space_before=12, space_after=2)
    
    # Add border bottom
    p_element = p._element
    pPr = p_element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)

def create_resume():
    doc = Document()
    
    # Header: Name
    header_name = doc.add_paragraph()
    header_name.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header_name.add_run("CHARLES S")
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = 'Calibri'
    set_spacing(header_name, 0, 5)

    # Contact Info
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = contact.add_run("+91-8971684901 | charlespinto204@gmail.com")
    run.font.size = Pt(10.5)
    set_spacing(contact, 0, 0)
    
    github = doc.add_paragraph()
    github.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = github.add_run("github.com/charles-codespace")
    run.font.size = Pt(10.5)
    set_spacing(github, 0, 10)

    # Projects Section
    add_section_header(doc, "Projects")
    
    # Project 1
    p1 = doc.add_paragraph()
    run = p1.add_run("DevOps Failure Prediction & Alerting System")
    run.bold = True
    run.font.size = Pt(12)
    set_spacing(p1, 5, 0)
    
    s1 = doc.add_paragraph()
    run = s1.add_run("Reva University, Bangalore")
    run.italic = True
    run.font.size = Pt(11)
    set_spacing(s1, 0, 5)
    
    bullets = [
        "Built a machine learning system to predict CI/CD failures using historical logs, improving reliability.",
        "Performed preprocessing and feature engineering to enhance model performance.",
        "Developed a real-time dashboard using Streamlit integrated with Elasticsearch.",
        "Containerized application using Docker for scalable deployment.",
        "Reduced response time and improved alerting efficiency."
    ]
    for b in bullets:
        bp = doc.add_paragraph(b, style='List Bullet')
        set_spacing(bp, 0, 2)
        bp.paragraph_format.left_indent = Pt(18)

    # Project 2
    p2 = doc.add_paragraph()
    run = p2.add_run("Full Stack IPsec Automation System")
    run.bold = True
    run.font.size = Pt(12)
    set_spacing(p2, 8, 0)
    
    s2 = doc.add_paragraph()
    run = s2.add_run("Reva University, Bangalore")
    run.italic = True
    run.font.size = Pt(11)
    set_spacing(s2, 0, 5)
    
    bullets2 = [
        "Built a centralized VPN automation system using IPsec (strongSwan).",
        "Developed backend APIs using FastAPI for secure policy management and API handing.",
        "Designed frontend using React.js for monitoring and visualization.",
        "Automated manual processes, reducing configuration errors.",
        "Improved efficiency and security of network operations."
    ]
    for b in bullets2:
        bp = doc.add_paragraph(b, style='List Bullet')
        set_spacing(bp, 0, 2)
        bp.paragraph_format.left_indent = Pt(18)

    # Skills Section
    add_section_header(doc, "Skills and Certifications")
    skills = [
        ("Programming Languages", "Java, Python, C++"),
        ("Tech Stack", "FastAPI, Flask, React.js, Docker, Streemlttreamlit, Scikit-learn, Elasticsearch"),
        ("Databases", "MySQL, Oracle, PostgreSQL, Elasticsearch"),
        ("Core Concepts", "Data Structures & Algorithms, DBMS, Computer Networks"),
        ("DevOps and Cloud", "CI/CD, AWS (Basics)"),
        ("Security", "IAM, RBAC, Authentication and Authorization, IPsec")
    ]
    for label, val in skills:
        p = doc.add_paragraph()
        run_l = p.add_run(f"• {label}: ")
        run_l.bold = True
        run_v = p.add_run(val)
        p.paragraph_format.left_indent = Pt(0)
        p.paragraph_format.first_line_indent = Pt(0)
        set_spacing(p, 0, 2)

    # Education Section
    add_section_header(doc, "Education")
    
    edu = doc.add_paragraph()
    run_univ = edu.add_run("Reva University, Bangalore")
    run_univ.bold = True
    run_univ.font.size = Pt(12)
    # Right align the date using a tab? Or just add it.
    # Simple way: just put it in the next paragraph or use tabs.
    edu.add_run("\t\t\t\t\t\t2023 – 2027")
    set_spacing(edu, 5, 0)
    
    edu_details = doc.add_paragraph()
    run = edu_details.add_run("B.Tech, Computer Science Enginneering\nCGPA: 8.2")
    run.font.size = Pt(11)
    set_spacing(edu_details, 0, 5)

    # Coursework Section
    add_section_header(doc, "Relevant Coursework")
    cw = doc.add_paragraph("• Data Structures & Algorithms, Operating Systems, Computer Networks, DBMS")
    set_spacing(cw, 5, 5)

    doc.save('resume.docx')

if __name__ == "__main__":
    create_resume()
